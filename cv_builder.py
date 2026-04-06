"""
فرصة | FURSA — مولّد السيرة الذاتية v4
ATS-Friendly CV — Arabic (Noto Naskh) & English
Outputs: PDF + DOCX
"""

import os, re
import arabic_reshaper
from bidi.algorithm import get_display
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

CV_DIR = "/tmp/cvs"
os.makedirs(CV_DIR, exist_ok=True)

# ── Noto Naskh Arabic font ──────────────────────────
ARABIC_FONT      = "Helvetica"
ARABIC_FONT_BOLD = "Helvetica-Bold"

for reg, bold in [
    ("./NotoNaskhArabic_400Regular.ttf",   "./NotoNaskhArabic_700Bold.ttf"),
    ("/app/NotoNaskhArabic_400Regular.ttf", "/app/NotoNaskhArabic_700Bold.ttf"),
]:
    try:
        if os.path.exists(reg) and os.path.exists(bold):
            pdfmetrics.registerFont(TTFont("NotoArabic",     reg))
            pdfmetrics.registerFont(TTFont("NotoArabicBold", bold))
            ARABIC_FONT      = "NotoArabic"
            ARABIC_FONT_BOLD = "NotoArabicBold"
            break
    except: pass

LATIN_FONT      = "Helvetica"
LATIN_FONT_BOLD = "Helvetica-Bold"
NAVY  = colors.HexColor("#0A1628")
TEAL  = colors.HexColor("#00A884")
GRAY  = colors.HexColor("#666666")
BLACK = colors.HexColor("#1A1A1A")
LGRAY = colors.HexColor("#E0E0E0")

# ── Arabic text processor ───────────────────────────
def ar(text):
    if not text: return ""
    try:
        return get_display(arabic_reshaper.reshape(str(text)))
    except:
        return str(text)

def has_arabic(text):
    return bool(re.search(r'[\u0600-\u06FF]', str(text)))

def smart_ar(text, is_ar_doc):
    """Only apply arabic reshaping if text contains Arabic chars."""
    if not text: return ""
    if is_ar_doc and has_arabic(str(text)):
        return ar(text)
    return str(text)

# ══════════════════════════════════════════════════════
#  CV STEPS
# ══════════════════════════════════════════════════════
CV_STEPS = {
    "en": [
        {"key": "lang",       "q": None},
        {"key": "name",       "q": "👤 What is your full name?"},
        {"key": "title",      "q": "💼 Your professional title:\n_(e.g. AI Engineer | Data Analyst)_"},
        {"key": "email",      "q": "📧 Your email address:"},
        {"key": "phone",      "q": "📱 Your phone number:"},
        {"key": "linkedin",   "q": "🔗 LinkedIn URL (or type 'skip'):"},
        {"key": "summary",    "q": "📝 Short professional summary:\n_(type 'ai' and I'll write it for you)_"},
        {"key": "skills",     "q": "🛠️ Key skills (comma-separated):\n_(e.g. Python, Power BI, SQL)_"},
        {"key": "education",  "q": "🎓 Education (one per line):\n`Degree | University | Year`\n_(e.g. MSc AI | University of Hail | 2024)_"},
        {"key": "experience", "q": "💼 Work experience — one job per message:\n`Job Title | Company | Start-End`\n`Achievement 1`\n\nType *done* when finished."},
        {"key": "courses",    "q": "📚 Courses & certifications (comma-separated or 'skip'):"},
        {"key": "languages",  "q": "🌍 Languages (e.g. Arabic Native, English Fluent):"},
    ],
    "ar": [
        {"key": "lang",       "q": None},
        {"key": "name",       "q": "👤 ما اسمك الكامل؟"},
        {"key": "title",      "q": "💼 ما مسماك الوظيفي؟\n_(مثال: مهندس ذكاء اصطناعي | محلل بيانات)_"},
        {"key": "email",      "q": "📧 عنوان إيميلك:"},
        {"key": "phone",      "q": "📱 رقم جوالك:"},
        {"key": "linkedin",   "q": "🔗 رابط LinkedIn (أو اكتب تخطي):"},
        {"key": "summary",    "q": "📝 نبذة مهنية مختصرة:\n_(اكتب ai وسأكتبها لك)_"},
        {"key": "skills",     "q": "🛠️ المهارات الرئيسية (مفصولة بفاصلة):"},
        {"key": "education",  "q": "🎓 التعليم (واحد في كل سطر):\n`الدرجة | الجامعة | السنة`\n_(مثال: ماجستير ذكاء اصطناعي | جامعة حائل | 2024)_"},
        {"key": "experience", "q": "💼 الخبرات — وظيفة لكل رسالة:\n`المسمى | الشركة | البداية-النهاية`\n`انجاز 1`\n\nاكتب انتهيت عند الإنهاء."},
        {"key": "courses",    "q": "📚 الدورات والشهادات (مفصولة بفاصلة أو تخطي):"},
        {"key": "languages",  "q": "🌍 اللغات (مثال: العربية لغة أم، الإنجليزية ممتاز):"},
    ]
}

# ══════════════════════════════════════════════════════
#  PDF GENERATOR
# ══════════════════════════════════════════════════════
def generate_cv_pdf(data: dict, chat_id: str) -> str:
    lang   = data.get("lang", "en")
    is_ar  = lang == "ar"
    fname  = f"{CV_DIR}/cv_{chat_id}_{lang}.pdf"
    f_reg  = ARABIC_FONT      if is_ar else LATIN_FONT
    f_bold = ARABIC_FONT_BOLD if is_ar else LATIN_FONT_BOLD
    align  = TA_RIGHT         if is_ar else TA_LEFT

    def t(text):
        return smart_ar(text, is_ar)

    doc = SimpleDocTemplate(fname, pagesize=A4,
        leftMargin=1.8*cm, rightMargin=1.8*cm,
        topMargin=1.5*cm, bottomMargin=1.5*cm)

    def ps(name, **kw): return ParagraphStyle(name, **kw)

    sName    = ps("N",  fontSize=20,   fontName=f_bold, textColor=NAVY,  alignment=align, spaceAfter=2,  leading=28)
    sTitle   = ps("T",  fontSize=11,   fontName=f_reg,  textColor=TEAL,  alignment=align, spaceAfter=3,  leading=18)
    sContact = ps("C",  fontSize=9,    fontName=LATIN_FONT, textColor=GRAY, alignment=align, spaceAfter=6, leading=12)
    sSec     = ps("S",  fontSize=10.5, fontName=f_bold, textColor=NAVY,  alignment=align, spaceBefore=10, spaceAfter=3, leading=18)
    sBody    = ps("B",  fontSize=10,   fontName=f_reg,  textColor=BLACK, alignment=align, leading=18, spaceAfter=2)
    sBullet  = ps("BU", fontSize=10,   fontName=f_reg,  textColor=BLACK, alignment=align, leading=18,
                  leftIndent=0 if is_ar else 10, rightIndent=12 if is_ar else 0, spaceAfter=1)
    sJT      = ps("JT", fontSize=10.5, fontName=f_bold, textColor=BLACK, alignment=align, spaceAfter=1, leading=18)
    # Education meta — always Latin font for university/year
    sJM_lat  = ps("JML", fontSize=9, fontName=LATIN_FONT, textColor=GRAY, alignment=align, spaceAfter=3, leading=12)

    story = []

    story.append(Paragraph(t(data.get("name", "")), sName))
    story.append(Paragraph(t(data.get("title", "")), sTitle))

    contacts = []
    if data.get("email"):  contacts.append(data["email"])
    if data.get("phone"):  contacts.append(data["phone"])
    if data.get("linkedin","").lower() not in ["skip","تخطي",""]:
        contacts.append(data["linkedin"])
    if contacts:
        story.append(Paragraph("  |  ".join(contacts), sContact))

    story.append(HRFlowable(width="100%", thickness=1.5, color=TEAL, spaceAfter=8))

    def section(en, ar_title):
        story.append(Paragraph(t(ar_title) if is_ar else en, sSec))
        story.append(HRFlowable(width="100%", thickness=0.4, color=LGRAY, spaceAfter=5))

    if data.get("summary"):
        section("ABOUT ME", "نبذة عني")
        story.append(Paragraph(t(data["summary"]), sBody))
        story.append(Spacer(1, 4))

    if data.get("experience"):
        section("EXPERIENCE", "الخبرات الوظيفية")
        for job in data["experience"]:
            lines = [l.strip() for l in job.strip().split("\n") if l.strip()]
            if not lines: continue
            header = lines[0].split("|")
            jt   = t(header[0].strip()) if len(header) > 0 else ""
            comp = header[1].strip()    if len(header) > 1 else ""
            date = header[2].strip()    if len(header) > 2 else ""
            story.append(Paragraph(jt, sJT))
            if comp or date:
                story.append(Paragraph(f"{comp}  •  {date}", sJM_lat))
            for line in lines[1:]:
                story.append(Paragraph(f"• {t(line)}", sBullet))
            story.append(Spacer(1, 5))

    if data.get("education"):
        section("EDUCATION", "التعليم")
        for edu in data["education"]:
            parts = edu.split("|")
            # Degree: may be Arabic — use t()
            deg = t(parts[0].strip()) if parts else t(edu)
            # University & year: always Latin
            uni = parts[1].strip() if len(parts) > 1 else ""
            yr  = parts[2].strip() if len(parts) > 2 else ""
            story.append(Paragraph(deg, sJT))
            if uni or yr:
                story.append(Paragraph(f"{uni}  •  {yr}", sJM_lat))
            story.append(Spacer(1, 3))

    if data.get("skills"):
        section("SKILLS", "المهارات")
        skills = [t(s.strip()) for s in data["skills"].split(",") if s.strip()]
        story.append(Paragraph("  •  ".join(skills), sBody))
        story.append(Spacer(1, 4))

    if data.get("courses","").lower() not in ["skip","تخطي",""]:
        section("COURSES & CERTIFICATIONS", "الدورات والشهادات")
        for c in data["courses"].split(","):
            c = c.strip()
            if c:
                story.append(Paragraph(f"• {t(c)}", sBullet))
        story.append(Spacer(1, 3))

    if data.get("languages"):
        section("LANGUAGES", "اللغات")
        story.append(Paragraph(t(data["languages"]), sBody))

    doc.build(story)
    return fname

# ══════════════════════════════════════════════════════
#  WORD GENERATOR
# ══════════════════════════════════════════════════════
def _set_color(run, hex_color):
    run.font.color.rgb = RGBColor(
        int(hex_color[0:2],16),
        int(hex_color[2:4],16),
        int(hex_color[4:6],16)
    )

def _add_hr(doc, color="00A884", thickness=12):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    return p

def generate_cv_docx(data: dict, chat_id: str) -> str:
    lang  = data.get("lang", "en")
    is_ar = lang == "ar"
    fname = f"{CV_DIR}/cv_{chat_id}_{lang}.docx"
    align = WD_ALIGN_PARAGRAPH.RIGHT if is_ar else WD_ALIGN_PARAGRAPH.LEFT
    ar_font  = "Arial"   # Word has built-in Arabic support
    lat_font = "Calibri"

    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    def add_para(text, bold=False, size=11, color=None, align_=None, font=None, space_after=4):
        p   = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)
        p.alignment = align_ or align
        run = p.add_run(text or "")
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = font or (ar_font if is_ar else lat_font)
        if color:
            _set_color(run, color)
        return p

    # Name
    add_para(data.get("name",""), bold=True, size=20, color="0A1628", space_after=2)
    # Title
    add_para(data.get("title",""), size=11, color="00A884", space_after=3)
    # Contact
    contacts = []
    if data.get("email"):  contacts.append(data["email"])
    if data.get("phone"):  contacts.append(data["phone"])
    if data.get("linkedin","").lower() not in ["skip","تخطي",""]:
        contacts.append(data["linkedin"])
    if contacts:
        add_para("  |  ".join(contacts), size=9, color="666666", font=lat_font, space_after=6)

    _add_hr(doc, color="00A884", thickness=18)

    def section(en, ar_title):
        add_para(ar_title if is_ar else en, bold=True, size=10, color="0A1628", space_after=2)
        _add_hr(doc, color="E0E0E0", thickness=6)

    def job_title(text):
        add_para(text, bold=True, size=10, color="1A1A1A", space_after=1)

    def job_meta(text):
        add_para(text, size=9, color="666666", font=lat_font, space_after=3)

    def body(text, space_after=3):
        add_para(text, size=10, space_after=space_after)

    def bullet(text):
        p   = doc.add_paragraph(style="List Bullet")
        p.alignment = align
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = ar_font if is_ar else lat_font

    if data.get("summary"):
        section("ABOUT ME", "نبذة عني")
        body(data["summary"])

    if data.get("experience"):
        section("EXPERIENCE", "الخبرات الوظيفية")
        for job in data["experience"]:
            lines = [l.strip() for l in job.strip().split("\n") if l.strip()]
            if not lines: continue
            header = lines[0].split("|")
            jt   = header[0].strip() if len(header) > 0 else ""
            comp = header[1].strip() if len(header) > 1 else ""
            date = header[2].strip() if len(header) > 2 else ""
            job_title(jt)
            if comp or date:
                job_meta(f"{comp}  •  {date}")
            for line in lines[1:]:
                bullet(line)

    if data.get("education"):
        section("EDUCATION", "التعليم")
        for edu in data["education"]:
            parts = edu.split("|")
            deg = parts[0].strip() if parts else edu
            uni = parts[1].strip() if len(parts) > 1 else ""
            yr  = parts[2].strip() if len(parts) > 2 else ""
            job_title(deg)
            if uni or yr:
                job_meta(f"{uni}  •  {yr}")

    if data.get("skills"):
        section("SKILLS", "المهارات")
        skills = [s.strip() for s in data["skills"].split(",") if s.strip()]
        body("  •  ".join(skills))

    if data.get("courses","").lower() not in ["skip","تخطي",""]:
        section("COURSES & CERTIFICATIONS", "الدورات والشهادات")
        for c in [c.strip() for c in data["courses"].split(",") if c.strip()]:
            bullet(c)

    if data.get("languages"):
        section("LANGUAGES", "اللغات")
        body(data["languages"])

    doc.save(fname)
    return fname

# ══════════════════════════════════════════════════════
#  AI SUMMARY
# ══════════════════════════════════════════════════════
def generate_ai_summary(data: dict, ai_client) -> str:
    lang  = data.get("lang", "en")
    is_ar = lang == "ar"
    try:
        r = ai_client.messages.create(
            model="claude-sonnet-4-20250514", max_tokens=200,
            messages=[{"role": "user", "content":
                f"Write a 2-3 sentence professional CV summary "
                f"{'in Arabic' if is_ar else 'in English'} for:\n"
                f"Title: {data.get('title','')}\nSkills: {data.get('skills','')}\n"
                f"Output the summary text only, no intro."}]
        )
        return r.content[0].text.strip()
    except:
        return ""
